#!/usr/bin/env python3
"""Minimal Markdown -> .docx converter for the Zero to VCAP posts.
Embeds local images, renders headings, bold/italic/code, links, lists,
blockquotes, and image captions. Tailored to the subset of Markdown used
in these blog drafts (no tables remain in this post)."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2])
BASE = SRC.parent

raw = SRC.read_text(encoding="utf-8")

# Strip YAML front matter, capture title
title = None
if raw.startswith("---"):
    end = raw.find("\n---", 3)
    fm = raw[3:end]
    m = re.search(r'title:\s*"?(.*?)"?\s*$', fm, re.MULTILINE)
    if m:
        title = m.group(1)
    raw = raw[end + 4:]

# Drop HTML comments
raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)

lines = raw.split("\n")

doc = Document()

INLINE = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*|\[[^\]]+\]\([^)]+\))"
)

def add_inline(p, text):
    """Add runs to paragraph p, honoring **bold**, *italic*, `code`, [link](url)."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif tok.startswith("*"):
            p.add_run(tok[1:-1]).italic = True
        elif tok.startswith("["):
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            r = p.add_run(lm.group(1)); r.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])

IMG = re.compile(r"!\[(.*?)\]\((.+?)\)")

if title:
    doc.add_heading(title, level=0)

i = 0
n = len(lines)
img_count = 0
while i < n:
    line = lines[i].rstrip()
    s = line.strip()

    if not s:
        i += 1
        continue

    # Horizontal rule
    if s == "---":
        p = doc.add_paragraph()
        r = p.add_run("• • •"); r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Image (own line)
    im = IMG.match(s)
    if im:
        alt, src = im.group(1), im.group(2)
        path = (BASE / src)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            p.add_run().add_picture(str(path), width=Inches(6.3))
            img_count += 1
        else:
            r = p.add_run(f"[missing image: {src}]"); r.font.color.rgb = RGBColor(0xC0, 0, 0)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(alt); cr.italic = True; cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        i += 1
        continue

    # Headings
    hm = re.match(r"(#{1,4})\s+(.*)", s)
    if hm:
        level = len(hm.group(1))
        doc.add_heading(hm.group(2), level=level)
        i += 1
        continue

    # Blockquote (may span lines)
    if s.startswith(">"):
        buf = []
        while i < n and lines[i].strip().startswith(">"):
            buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
            i += 1
        p = doc.add_paragraph(style="Quote")
        add_inline(p, " ".join(x.strip() for x in buf if x.strip()))
        continue

    # Bullet list
    if re.match(r"[-*]\s+", s):
        while i < n and re.match(r"\s*[-*]\s+", lines[i]):
            item = re.sub(r"^\s*[-*]\s+", "", lines[i])
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, item.strip())
            i += 1
        continue

    # Paragraph (gather until blank / block boundary)
    buf = [line]
    i += 1
    while i < n:
        nxt = lines[i].strip()
        if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith(">")
                or IMG.match(nxt) or re.match(r"[-*]\s+", nxt)):
            break
        buf.append(lines[i])
        i += 1
    p = doc.add_paragraph()
    add_inline(p, " ".join(x.strip() for x in buf))

doc.save(str(OUT))
print(f"Wrote {OUT}  ({img_count} images embedded)")
