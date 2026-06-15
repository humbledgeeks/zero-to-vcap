#!/usr/bin/env python3
"""
Build VCF-9.1-WorkloadDomain-Outline.docx as a faithful render of the Part 2 blog
post (blog-posts/zero-to-vcap-vcf91-workload-domain.md): intro + URLs + headings +
tables + tip/gotcha callouts. Screenshot placeholders are HTML comments, so the doc
renders as a clean text outline until real images are added.

Mirrors _build_doc.py (Part 1). Run from the repo root:

    .venv/bin/python vcf-deployment-blog/_build_wld_doc.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parent.parent
BLOG = REPO / "blog-posts" / "zero-to-vcap-vcf91-workload-domain.md"
BLOG_DIR = BLOG.parent
DOCX_OUT = Path(__file__).resolve().parent / "VCF-9.1-WorkloadDomain-Outline.docx"

IMG_LINE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
# Inline: link | bold | inline-code
INLINE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text):
    """Render inline markdown (links, bold, code) into a paragraph."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:  # link [text](url)
            add_hyperlink(paragraph, m.group(1), m.group(2))
        elif m.group(3) is not None:  # **bold**
            paragraph.add_run(m.group(3)).bold = True
        elif m.group(4) is not None:  # `code`
            r = paragraph.add_run(m.group(4))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main():
    lines = BLOG.read_text().splitlines()
    doc = Document()
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # HTML comment -> skip
        if stripped.startswith("<!--"):
            i += 1
            continue

        # blank -> skip (block spacing handled per-block)
        if not stripped:
            i += 1
            continue

        # horizontal rule -> small spacer
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # image
        m = IMG_LINE.match(line)
        if m:
            img_path = (BLOG_DIR / m.group("path")).resolve()
            fname = img_path.name
            lbl = doc.add_paragraph()
            run = lbl.add_run(fname)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.exists():
                pic.add_run().add_picture(str(img_path), width=Inches(6.3))
            else:
                pic.add_run(f"[missing image: {fname}]").italic = True
            i += 1
            continue

        # heading
        m = HEADING.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc_level = {1: 0, 2: 1, 3: 2}.get(level, 3)
            doc.add_heading(text, level=doc_level)
            i += 1
            continue

        # blockquote (collect consecutive)
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, " ".join(buf))
            for r in p.runs:
                r.italic = True
            continue

        # table (collect consecutive | rows)
        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            cells = [
                [c.strip() for c in r.strip("|").split("|")]
                for r in rows
                if not re.match(r"^\|[\s:|-]+\|?$", r)
            ]
            if cells:
                tbl = doc.add_table(rows=len(cells), cols=len(cells[0]))
                tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(cells):
                    for ci, val in enumerate(row):
                        cell = tbl.rows[ri].cells[ci]
                        cell.text = ""
                        add_inline(cell.paragraphs[0], val)
            continue

        # bullet list (collect consecutive). Handles "- [ ] " checklist items too.
        if re.match(r"^[-*]\s+", stripped):
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                content = re.sub(r"^[-*]\s+", "", lines[i].strip())
                content = re.sub(r"^\[[ xX]\]\s*", "", content)  # strip checkbox marker
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, content)
                i += 1
            continue

        # paragraph (collect consecutive plain lines)
        buf = []
        while i < n:
            s = lines[i].strip()
            if (not s or s.startswith(("#", ">", "|", "<!--"))
                    or s == "---" or IMG_LINE.match(lines[i])
                    or re.match(r"^[-*]\s+", s)):
                break
            buf.append(s)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(buf))

    doc.save(DOCX_OUT)
    print(f"Built {DOCX_OUT.name} from {BLOG.name}")


if __name__ == "__main__":
    main()
